from __future__ import annotations

import re
from pathlib import Path

import fitz
import pdfplumber
from docx import Document

from app.utils.text import normalize_whitespace


class ResumeParsingError(Exception):
    """Raised when resume text extraction fails."""


def _score_text_quality(text: str) -> float:
    if not text:
        return 0.0

    length_score = min(len(text) / 2500.0, 1.0)
    words = re.findall(r"[A-Za-z][A-Za-z0-9+#.\-/]*", text)
    word_score = min(len(words) / 450.0, 1.0)

    printable = sum(ch.isprintable() for ch in text)
    printable_score = printable / max(len(text), 1)

    alpha = sum(ch.isalpha() for ch in text)
    symbol_penalty = 1.0 - min((len(text) - alpha) / max(len(text), 1), 0.6)

    return (0.40 * length_score) + (0.35 * word_score) + (0.15 * printable_score) + (0.10 * symbol_penalty)


def _extract_text_pdf_pymupdf(file_path: Path) -> str:
    chunks: list[str] = []
    with fitz.open(file_path) as doc:
        for page in doc:
            chunks.append(page.get_text("text"))
    return "\n".join(chunks)


def _extract_text_pdf_pdfplumber(file_path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def _extract_text_docx(file_path: Path) -> str:
    doc = Document(file_path)
    chunks: list[str] = []

    for para in doc.paragraphs:
        chunks.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells if cell.text)
            if row_text.strip():
                chunks.append(row_text)

    return "\n".join(chunks)


def extract_resume_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        parser_outputs: list[tuple[str, str, float]] = []

        for _name, parser in (
            ("pymupdf", _extract_text_pdf_pymupdf),
            ("pdfplumber", _extract_text_pdf_pdfplumber),
        ):
            try:
                text = parser(file_path)
            except Exception:
                text = ""
            parser_outputs.append((_name, text, _score_text_quality(text)))

        _, best_text, best_score = max(parser_outputs, key=lambda item: item[2])
        if best_score <= 0 or not best_text.strip():
            raise ResumeParsingError("Unable to extract readable text from PDF resume.")
        return normalize_whitespace(best_text)

    if suffix == ".docx":
        try:
            text = _extract_text_docx(file_path)
        except Exception as exc:
            raise ResumeParsingError(f"Unable to parse DOCX resume: {exc}") from exc
        if not text.strip():
            raise ResumeParsingError("Extracted empty text from DOCX resume.")
        return normalize_whitespace(text)

    raise ResumeParsingError("Unsupported resume type. Please upload PDF or DOCX.")
